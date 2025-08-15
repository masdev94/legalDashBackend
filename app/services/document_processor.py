import os
import re
import uuid
import io
from datetime import datetime
from typing import List, Dict, Any, Optional
from pathlib import Path
import PyPDF2
from docx import Document as DocxDocument
from loguru import logger

from app.models.document import (
    Document, DocumentMetadata, AgreementType, Jurisdiction, 
    Industry, Geography, AIInsights, RiskLevel, ComplianceStatus
)
from app.services.ai_analyzer import AIAnalyzer

class DocumentProcessor:
    
    def __init__(self):
        self.agreement_patterns = {
            AgreementType.NDA: r'\b(?:NDA|Non-Disclosure|Non Disclosure|Confidentiality)\b',
            AgreementType.MSA: r'\b(?:MSA|Master Service|Master Services)\b',
            AgreementType.FRANCHISE: r'\b(?:Franchise|Franchising)\b',
            AgreementType.EMPLOYMENT: r'\b(?:Employment|Employment Contract|Work Contract)\b',
            AgreementType.LEASE: r'\b(?:Lease|Leasing|Rental)\b',
            AgreementType.PARTNERSHIP: r'\b(?:Partnership|Joint Venture|Collaboration)\b',
            AgreementType.SUPPLY: r'\b(?:Supply|Procurement|Purchase)\b',
            AgreementType.SERVICE: r'\b(?:Service|Consulting|Advisory)\b',
            AgreementType.LICENSING: r'\b(?:License|Licensing|IP License)\b',
        }
        
        self.jurisdiction_patterns = {
            Jurisdiction.UAE: r'\b(?:UAE|United Arab Emirates|Dubai|Abu Dhabi|Sharjah)\b',
            Jurisdiction.UK: r'\b(?:UK|United Kingdom|England|Wales|Scotland|Northern Ireland)\b',
            Jurisdiction.USA: r'\b(?:USA|United States|US|America)\b',
            Jurisdiction.DELAWARE: r'\b(?:Delaware|DE)\b',
            Jurisdiction.SINGAPORE: r'\b(?:Singapore|SG)\b',
            Jurisdiction.HONG_KONG: r'\b(?:Hong Kong|HK)\b',
            Jurisdiction.GERMANY: r'\b(?:Germany|DE|Deutschland)\b',
            Jurisdiction.FRANCE: r'\b(?:France|FR|République française)\b',
        }
        
        self.industry_patterns = {
            Industry.TECHNOLOGY: r'\b(?:Technology|Tech|Software|IT|Information Technology|Digital)\b',
            Industry.HEALTHCARE: r'\b(?:Healthcare|Health|Medical|Pharmaceutical|Biotech)\b',
            Industry.FINANCE: r'\b(?:Finance|Financial|Banking|Investment|Insurance)\b',
            Industry.OIL_GAS: r'\b(?:Oil|Gas|Petroleum|Energy|Hydrocarbon)\b',
            Industry.REAL_ESTATE: r'\b(?:Real Estate|Property|Realty|Construction)\b',
            Industry.MANUFACTURING: r'\b(?:Manufacturing|Manufacture|Production|Factory)\b',
            Industry.RETAIL: r'\b(?:Retail|Commerce|Trading|Merchandise)\b',
            Industry.CONSULTING: r'\b(?:Consulting|Consultancy|Advisory|Advisory Services)\b',
        }
        
        self.geography_patterns = {
            Geography.MIDDLE_EAST: r'\b(?:Middle East|Gulf|GCC|Arab|Arabian|Persian Gulf)\b',
            Geography.EUROPE: r'\b(?:Europe|European|EU|Eurozone)\b',
            Geography.NORTH_AMERICA: r'\b(?:North America|North American|NAFTA)\b',
            Geography.ASIA_PACIFIC: r'\b(?:Asia|Asian|Pacific|APAC|Asia-Pacific)\b',
            Geography.AFRICA: r'\b(?:Africa|African|Sub-Saharan|North Africa)\b',
            Geography.SOUTH_AMERICA: r'\b(?:South America|Latin America|LATAM)\b',
        }
        
        self.governing_law_patterns = [
            r'\b(?:governed by|subject to|in accordance with)\s+([A-Za-z\s]+?)(?:law|jurisdiction|legal system)\b',
            r'\b(?:applicable law|governing law)\s*[:\-]?\s*([A-Za-z\s]+?)(?:law|jurisdiction|legal system)\b',
            r'\b(?:This agreement|This contract)\s+(?:shall be|is)\s+(?:governed by|subject to)\s+([A-Za-z\s]+?)(?:law|jurisdiction|legal system)\b',
        ]
        
        self.party_patterns = [
            r'\b(?:between|by and between)\s+([A-Za-z\s&.,]+?)\s+and\s+([A-Za-z\s&.,]+?)(?:\s|$)',
            r'\b(?:Party A|First Party)\s*[:\-]?\s*([A-Za-z\s&.,]+?)(?:\s|$)',
            r'\b(?:Party B|Second Party)\s*[:\-]?\s*([A-Za-z\s&.,]+?)(?:\s|$)',
        ]
        
        self.value_patterns = [
            r'\b(?:value|amount|consideration|price)\s*[:\-]?\s*([$€£¥]?\s*\d{1,3}(?:,\d{3})*(?:\.\d{2})?)\s*([A-Z]{3})?\b',
            r'\b([$€£¥]?\s*\d{1,3}(?:,\d{3})*(?:\.\d{2})?)\s*([A-Z]{3})?\s*(?:USD|EUR|GBP|AED|SAR)\b',
        ]
        
        self.date_patterns = [
            r'\b(?:effective date|commencement date)\s*[:\-]?\s*(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\b',
            r'\b(?:expiration date|termination date|end date)\s*[:\-]?\s*(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\b',
        ]

        self.ai_analyzer = AIAnalyzer()
        logger.info("Document processor initialized with AI analyzer")

    def process_document(self, file_path: str, file_content: bytes) -> Document:
        try:
            logger.info(f"=== STARTING DOCUMENT PROCESSING ===")
            logger.info(f"Processing file: {file_path}")
            logger.info(f"File size: {len(file_content)} bytes")
            logger.info(f"File content starts with: {file_content[:20]}")
            
            extracted_text = self._extract_text(file_path, file_content)
            logger.info(f"Extracted text length: {len(extracted_text)}")
            logger.info(f"First 100 chars of text: {extracted_text[:100]}")
            
            metadata = self._extract_metadata(file_path, extracted_text)
            logger.info(f"Final metadata: {metadata}")
            
            # Generate AI insights for the document during upload
            ai_insights = self.ai_analyzer._generate_basic_ai_insights(Document(
                id="temp",
                metadata=metadata,
                content=extracted_text,
                extracted_text=extracted_text,
                created_at=datetime.now(),
                updated_at=datetime.now()
            ))
            logger.info(f"AI insights generated: {ai_insights is not None}")
            
            document = Document(
                id=str(uuid.uuid4()),
                metadata=metadata,
                content=file_content,
                extracted_text=extracted_text,
                raw_bytes=file_content,
                ai_insights=ai_insights,
                confidence_score=ai_insights.confidence_score if ai_insights else None,
                created_at=datetime.now(),
                updated_at=datetime.now()
            )
            
            logger.info(f"Document created with ID: {document.id}")
            logger.info(f"Content type: {type(document.content)}")
            logger.info(f"Content length: {len(document.content) if isinstance(document.content, bytes) else 'N/A'}")
            logger.info(f"=== DOCUMENT PROCESSING COMPLETE ===")
            return document
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            logger.error(f"Exception type: {type(e)}")
            import traceback
            logger.error(f"Traceback: {traceback.format_exc()}")
            raise

    def _extract_text(self, file_path: str, file_content: bytes) -> str:
        try:
            file_extension = self._detect_file_type(file_content)
            logger.info(f"Detected file type: {file_extension}")
            
            if file_extension == '.pdf':
                return self._extract_pdf_text(file_content)
            elif file_extension == '.docx':
                return self._extract_docx_text(file_content)
            elif file_extension == '.doc':
                return self._extract_doc_text(file_content)
            else:
                logger.warning(f"Unsupported file type: {file_extension}")
                return ""
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return ""

    def _detect_file_type(self, file_content: bytes) -> str:
        try:
            if file_content.startswith(b'%PDF'):
                return '.pdf'
            elif file_content.startswith(b'PK'):
                return '.docx'
            elif file_content.startswith(b'\xd0\xcf\x11\xe0'):
                return '.doc'
            else:
                return '.pdf'
        except Exception:
            return '.pdf'

    def _extract_pdf_text(self, file_content: bytes) -> str:
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting PDF text: {str(e)}")
            return ""

    def _extract_docx_text(self, file_content: bytes) -> str:
        try:
            doc = DocxDocument(io.BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            return ""

    def _extract_doc_text(self, file_content: bytes) -> str:
        try:
            return "DOC file content extracted"
        except Exception as e:
            logger.error(f"Error extracting DOC text: {str(e)}")
            return ""

    def _extract_metadata(self, file_path: str, extracted_text: str) -> DocumentMetadata:
        try:
            file_path_obj = Path(file_path)
            file_size = len(extracted_text.encode('utf-8')) if extracted_text else 0
            
            metadata = DocumentMetadata(
                filename=file_path_obj.name,
                file_size=file_size,
                file_type=file_path_obj.suffix.lower(),
                upload_date=datetime.now()
            )
            
            if extracted_text and len(extracted_text.strip()) > 0:
                metadata.agreement_type = self._extract_agreement_type(extracted_text)
                
                metadata.jurisdiction = self._extract_jurisdiction(extracted_text)
                
                metadata.industry = self._extract_industry(extracted_text)
                
                metadata.geography = self._extract_geography(extracted_text)
                
                metadata.governing_law = self._extract_governing_law(extracted_text)
                metadata.parties = self._extract_parties(extracted_text)
                metadata.effective_date = self._extract_effective_date(extracted_text)
                metadata.expiration_date = self._extract_expiration_date(extracted_text)
                metadata.value, metadata.currency = self._extract_value(extracted_text)
            
            return metadata
            
        except Exception as e:
            logger.error(f"Error extracting metadata from {file_path}: {str(e)}")
            return DocumentMetadata(
                filename=Path(file_path).name,
                file_size=0,
                file_type=Path(file_path).suffix.lower(),
                upload_date=datetime.now()
            )

    def _extract_agreement_type(self, text: str) -> Optional[AgreementType]:
        for agreement_type, pattern in self.agreement_patterns.items():
            if re.search(pattern, text, re.IGNORECASE):
                return agreement_type
        return None

    def _extract_jurisdiction(self, text: str) -> Optional[Jurisdiction]:
        for jurisdiction, pattern in self.jurisdiction_patterns.items():
            if re.search(pattern, text, re.IGNORECASE):
                return jurisdiction
        return None

    def _extract_industry(self, text: str) -> Optional[Industry]:
        for industry, pattern in self.industry_patterns.items():
            if re.search(pattern, text, re.IGNORECASE):
                return industry
        return None

    def _extract_geography(self, text: str) -> Optional[Geography]:
        for geography, pattern in self.geography_patterns.items():
            if re.search(pattern, text, re.IGNORECASE):
                return geography
        return None

    def _extract_governing_law(self, text: str) -> Optional[str]:
        for pattern in self.governing_law_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        return None

    def _extract_parties(self, text: str) -> List[str]:
        parties = []
        for pattern in self.party_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                if isinstance(match, tuple):
                    parties.extend([party.strip() for party in match if party.strip()])
                else:
                    parties.append(match.strip())
        return list(set(parties))

    def _extract_effective_date(self, text: str) -> Optional[datetime]:
        for pattern in self.date_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                try:
                    date_str = match.group(1)
                    return datetime.now()
                except Exception:
                    continue
        return None

    def _extract_expiration_date(self, text: str) -> Optional[datetime]:
        return None

    def _extract_value(self, text: str) -> tuple[Optional[float], Optional[str]]:
        for pattern in self.value_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                try:
                    value_str = match.group(1).replace(',', '')
                    currency = match.group(2) if len(match.groups()) > 1 else None
                    value = float(value_str.replace('$', '').replace('€', '').replace('£', '').replace('¥', ''))
                    return value, currency
                except Exception:
                    continue
        return None, None
    def _generate_document_ai_insights(self, extracted_text: str, metadata: DocumentMetadata) -> AIInsights:
        try:
            if not extracted_text or len(extracted_text.strip()) < 50:
                logger.warning("Text too short for meaningful AI analysis")
                return AIInsights(
                    summary="📄 **Document Analysis**\n\nDocument text too short for comprehensive analysis. Please ensure the document contains sufficient content for AI processing.",
                    confidence_score=0.1
                )
            
            summary = self._generate_document_summary(extracted_text, metadata)
            key_terms = self._extract_key_terms(extracted_text, metadata.filename)
            risk_level, risk_factors = self._assess_document_risk(extracted_text)
            compliance_status, compliance_reqs = self._assess_compliance_status(extracted_text)
            business_impact = self._analyze_business_impact(extracted_text)
            recommendations = self._generate_recommendations(risk_level, compliance_status, business_impact)
            extracted_entities = self._extract_entities(extracted_text)
            sentiment = self._analyze_sentiment(extracted_text)
            complexity_score = self._calculate_complexity_score(extracted_text)
            confidence_score = self._calculate_confidence_score(extracted_text, key_terms)
            
            ai_insights = AIInsights(
                summary=summary,
                key_terms=key_terms,
                risk_assessment=risk_level,
                compliance_status=compliance_status,
                business_impact=business_impact,
                recommendations=recommendations,
                confidence_score=confidence_score,
                extracted_entities=extracted_entities,
                sentiment_analysis=sentiment,
                complexity_score=complexity_score
            )
            
            logger.info("Fresh AI insights generated successfully with confidence: %s", confidence_score)
            return ai_insights
            
        except Exception as e:
            logger.error("Error generating document AI insights: %s", str(e))
            return AIInsights(
                summary="📄 **Document Analysis**\n\nAI analysis encountered an error. Please try again or contact support.",
                confidence_score=0.1
            )

    def _generate_document_summary(self, text: str, metadata: DocumentMetadata) -> str:
        """Generate a concise summary of the document"""
        try:
            # Use AI analyzer for summary generation
            summary = self.ai_analyzer.generate_document_summary(text, metadata)
            return summary
        except Exception as e:
            logger.error(f"Error generating document summary: {str(e)}")
            # Fallback to basic summary
            return f"Legal document ({metadata.agreement_type or 'Unknown type'}) from {metadata.jurisdiction or 'Unknown jurisdiction'} jurisdiction. Document contains {len(text.split())} words of legal content."

    def _extract_key_terms(self, text: str, filename: str) -> List[str]:
        try:
            terms = set()
            
            legal_terms = [
                'contract', 'agreement', 'terms', 'conditions', 'clause', 'section',
                'party', 'parties', 'obligation', 'liability', 'breach', 'termination',
                'renewal', 'amendment', 'waiver', 'indemnification', 'governing law'
            ]
            
            financial_terms = [
                'payment', 'fee', 'cost', 'price', 'amount', 'revenue', 'profit',
                'loss', 'budget', 'investment', 'funding', 'currency', 'dollar'
            ]
            
            business_terms = [
                'partnership', 'collaboration', 'joint venture', 'merger', 'acquisition',
                'market', 'competition', 'strategy', 'growth', 'expansion'
            ]
            
            compliance_terms = [
                'compliance', 'regulation', 'legal', 'statute', 'requirement',
                'license', 'permit', 'certification', 'accreditation', 'audit'
            ]
            
            for term in legal_terms:
                if term in text.lower():
                    terms.add(term.title())
            
            for term in financial_terms:
                if term in text.lower():
                    terms.add(term.title())
            
            for term in business_terms:
                if term in text.lower():
                    terms.add(term.title())
            
            for term in compliance_terms:
                if term in text.lower():
                    terms.add(term.title())
            
            if not terms:
                terms.add("Legal Document")
                terms.add("Agreement")
            
            return list(terms)[:8]
            
        except Exception as e:
            logger.error(f"Error extracting key terms: {str(e)}")
            return ["Legal", "Document"]

    def _assess_document_risk(self, text: str) -> tuple[RiskLevel, List[str]]:
        try:
            high_risk_keywords = ['breach', 'termination', 'penalty', 'damages', 'liability', 'indemnification']
            medium_risk_keywords = ['obligation', 'requirement', 'condition', 'deadline', 'performance']
            
            high_risk_count = sum(1 for keyword in high_risk_keywords if keyword.lower() in text.lower())
            medium_risk_count = sum(1 for keyword in medium_risk_keywords if keyword.lower() in text.lower())
            
            risk_level = RiskLevel.HIGH if high_risk_count >= 3 else RiskLevel.MEDIUM if medium_risk_count >= 2 or high_risk_count >= 1 else RiskLevel.LOW
            risk_factors = [f"High risk: {keyword}" for keyword in high_risk_keywords if keyword.lower() in text.lower()] + \
                           [f"Medium risk: {keyword}" for keyword in medium_risk_keywords if keyword.lower() in text.lower()]
            return risk_level, risk_factors
            
        except Exception as e:
            logger.error(f"Error assessing document risk: {str(e)}")
            return RiskLevel.MEDIUM, ["Risk assessment failed"]

    def _assess_compliance_status(self, text: str) -> tuple[ComplianceStatus, List[str]]:
        try:
            compliance_keywords = ['compliance', 'regulation', 'requirement', 'standard', 'guideline']
            non_compliance_keywords = ['violation', 'breach', 'non-compliance', 'penalty']
            
            compliance_count = sum(1 for keyword in compliance_keywords if keyword.lower() in text.lower())
            non_compliance_count = sum(1 for keyword in non_compliance_keywords if keyword.lower() in text.lower())
            
            compliance_status = ComplianceStatus.NON_COMPLIANT if non_compliance_count > compliance_count else ComplianceStatus.COMPLIANT if compliance_count > 0 else ComplianceStatus.PENDING_REVIEW
            compliance_reqs = [f"Compliance requirement: {keyword}" for keyword in compliance_keywords if keyword.lower() in text.lower()] + \
                             [f"Non-compliance: {keyword}" for keyword in non_compliance_keywords if keyword.lower() in text.lower()]
            return compliance_status, compliance_reqs
            
        except Exception as e:
            logger.error(f"Error assessing compliance status: {str(e)}")
            return ComplianceStatus.PENDING_REVIEW, ["Compliance assessment failed"]

    def _analyze_business_impact(self, text: str) -> str:
        try:
            financial_patterns = [
                r'\b(?:revenue|profit|loss|cost|expense|budget|investment|funding)\b',
                r'\b(?:dollar|euro|pound|currency|amount|value|price|fee)\b'
            ]
            
            operational_patterns = [
                r'\b(?:operation|process|efficiency|productivity|workflow|procedure)\b',
                r'\b(?:resource|staff|personnel|equipment|infrastructure)\b'
            ]
            
            strategic_patterns = [
                r'\b(?:market|competition|strategy|growth|expansion|partnership)\b',
                r'\b(?:merger|acquisition|joint venture|alliance|collaboration)\b'
            ]
            
            regulatory_patterns = [
                r'\b(?:regulation|compliance|legal|statute|law|requirement)\b',
                r'\b(?:license|permit|certification|accreditation)\b'
            ]
            
            impact_scores = {'financial': 0, 'operational': 0, 'strategic': 0, 'regulatory': 0}
            
            for pattern in financial_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                impact_scores['financial'] += len(matches)
            
            for pattern in operational_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                impact_scores['operational'] += len(matches)
            
            for pattern in strategic_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                impact_scores['strategic'] += len(matches)
            
            for pattern in regulatory_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                impact_scores['regulatory'] += len(matches)
            
            max_impact = max(impact_scores.values())
            if max_impact == 0:
                return "Standard business operations"
            
            impact_areas = []
            if impact_scores['financial'] > 2:
                impact_areas.append("Financial")
            if impact_scores['operational'] > 2:
                impact_areas.append("Operational")
            if impact_scores['strategic'] > 2:
                impact_areas.append("Strategic")
            if impact_scores['regulatory'] > 2:
                impact_areas.append("Regulatory")
            
            if not impact_areas:
                if max_impact > 1:
                    impact_areas.append("Moderate")
                else:
                    impact_areas.append("Limited")
            
            return f"{', '.join(impact_areas)} - {self._get_impact_description(impact_areas, max_impact)}"
            
        except Exception as e:
            logger.error(f"Error analyzing business impact: {str(e)}")
            return "Standard business impact"

    def _get_impact_description(self, impact_areas: List[str], max_score: int) -> str:
        if "Financial" in impact_areas and max_score > 3:
            return "Significant financial implications"
        elif "Strategic" in impact_areas and max_score > 3:
            return "Strategic business decisions required"
        elif "Regulatory" in impact_areas and max_score > 3:
            return "Regulatory compliance critical"
        elif "Operational" in impact_areas and max_score > 3:
            return "Operational processes affected"
        elif max_score > 2:
            return "Moderate business impact"
        else:
            return "Standard business impact"

    def _generate_recommendations(self, risk_level: RiskLevel, compliance_status: ComplianceStatus, business_impact: str) -> List[str]:
        try:
            recommendations = []
            
            if risk_level == RiskLevel.HIGH:
                recommendations.append("🔴 Immediate legal review required")
                recommendations.append("⚠️ Implement risk mitigation strategies")
                recommendations.append("📊 Monitor for potential breaches")
            elif risk_level == RiskLevel.MEDIUM:
                recommendations.append("🟡 Schedule legal review within 30 days")
                recommendations.append("📈 Review risk factors quarterly")
            elif risk_level == RiskLevel.LOW:
                recommendations.append("🟢 Standard review process")
            
            if compliance_status == ComplianceStatus.REQUIRES_ACTION:
                recommendations.append("🚨 Urgent compliance review needed")
                recommendations.append("📋 Implement compliance monitoring")
            elif compliance_status == ComplianceStatus.PENDING_REVIEW:
                recommendations.append("⏰ Schedule compliance review")
            elif compliance_status == ComplianceStatus.COMPLIANT:
                recommendations.append("✅ Maintain current compliance standards")
            
            if "Financial" in business_impact:
                recommendations.append("💰 Financial impact assessment required")
                recommendations.append("📊 Monitor financial metrics closely")
            if "Strategic" in business_impact:
                recommendations.append("🎯 Executive review recommended")
                recommendations.append("📈 Strategic planning session needed")
            if "Regulatory" in business_impact:
                recommendations.append("⚖️ Legal compliance verification")
                recommendations.append("📋 Regulatory monitoring setup")
            
            if not recommendations:
                recommendations.append("📋 Standard review process")
                recommendations.append("👀 Monitor for changes")
            
            return recommendations[:6]
            
        except Exception as e:
            logger.error(f"Error generating recommendations: {str(e)}")
            return ["Review document manually", "Contact legal team if needed"]

    def _extract_entities(self, text: str) -> Dict[str, Any]:
        try:
            entities = {
                'organizations': [],
                'locations': [],
                'dates': [],
                'amounts': [],
                'legal_terms': []
            }
            
            org_patterns = [
                r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s+(?:Inc|Corp|LLC|Ltd|Company|Corporation)\b',
                r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s+(?:Partners|Associates|Group)\b'
            ]
            
            for pattern in org_patterns:
                matches = re.findall(pattern, text)
                entities['organizations'].extend(matches)
            
            return entities
            
        except Exception as e:
            logger.error(f"Error extracting entities: {str(e)}")
            return {'organizations': [], 'locations': [], 'dates': [], 'amounts': [], 'legal_terms': []}

    def _analyze_sentiment(self, text: str) -> str:
        try:
            positive_words = ['benefit', 'advantage', 'opportunity', 'success', 'growth', 'positive']
            negative_words = ['risk', 'liability', 'penalty', 'breach', 'termination', 'damage']
            
            positive_count = sum(1 for word in positive_words if word.lower() in text.lower())
            negative_count = sum(1 for word in negative_words if word.lower() in text.lower())
            
            if positive_count > negative_count:
                return "Positive - Favorable terms and conditions"
            elif negative_count > positive_count:
                return "Negative - Contains risk factors and liabilities"
            else:
                return "Neutral - Balanced legal terms"
                
        except Exception as e:
            logger.error(f"Error analyzing sentiment: {str(e)}")
            return "Sentiment analysis unavailable"

    def _calculate_complexity_score(self, text: str) -> float:
        try:
            score = 0.5
            
            if len(text) > 2000:
                score += 0.3
            elif len(text) > 1000:
                score += 0.2
            
            complex_terms = ['indemnification', 'jurisdiction', 'arbitration', 'mediation', 'governing law']
            complex_count = sum(1 for term in complex_terms if term.lower() in text.lower())
            
            if complex_count > 3:
                score += 0.2
            elif complex_count > 1:
                score += 0.1
            
            return min(score, 1.0)
            
        except Exception as e:
            logger.error(f"Error calculating complexity score: {str(e)}")
            return 0.5

    def _calculate_confidence_score(self, text: str, key_terms: List[str]) -> float:
        try:
            score = 0.5
            
            if len(text) > 1000:
                score += 0.2
            elif len(text) > 500:
                score += 0.1
            
            if len(key_terms) > 5:
                score += 0.15
            elif len(key_terms) > 2:
                score += 0.1
            
            if any(term in text for term in ['contract', 'agreement', 'terms', 'conditions']):
                score += 0.1
            
            if any(term in text for term in ['party', 'parties', 'obligation', 'liability']):
                score += 0.05
            
            return min(score, 1.0)
            
        except Exception as e:
            logger.error(f"Error calculating confidence score: {str(e)}")
            return 0.5
